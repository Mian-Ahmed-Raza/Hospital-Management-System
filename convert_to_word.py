"""
Convert PROJECT_DOCUMENTATION.md to Word document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_to_word(md_file, output_file):
    """Convert Markdown file to Word document"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            # H1 - Title
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.color.rgb = RGBColor(52, 152, 219)
            
        elif line.startswith('## '):
            # H2
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(52, 152, 219)
            
        elif line.startswith('### '):
            # H3
            p = doc.add_heading(line[4:], level=3)
            p.runs[0].font.color.rgb = RGBColor(46, 204, 113)
            
        elif line.startswith('#### '):
            # H4
            p = doc.add_heading(line[5:], level=4)
            
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
            
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            
            # Add code as paragraph with monospace font
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Check if it's a checklist item
            if text.startswith('[ ] '):
                text = '☐ ' + text[4:]
            elif text.startswith('[x] ') or text.startswith('[X] '):
                text = '☑ ' + text[4:]
            
            p = doc.add_paragraph(text, style='List Bullet')
            format_inline_markdown(p)
            
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            format_inline_markdown(p)
            
        # Block quotes
        elif line.startswith('> '):
            text = line[2:]
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.italic = True
            
        # Tables (simple detection)
        elif '|' in line and line.count('|') > 2:
            # Skip table for now (complex parsing needed)
            while i < len(lines) and '|' in lines[i]:
                i += 1
            continue
            
        # Bold/Strong text with **text**
        else:
            # Regular paragraph
            if line.strip():
                p = doc.add_paragraph(line)
                format_inline_markdown(p)
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✅ Document saved as: {output_file}")

def format_inline_markdown(paragraph):
    """Format inline markdown (bold, italic, code)"""
    text = paragraph.text
    paragraph.clear()
    
    # Pattern for inline formatting
    # **bold**, *italic*, `code`
    parts = []
    current = ""
    i = 0
    
    while i < len(text):
        # Bold **text**
        if i < len(text) - 1 and text[i:i+2] == '**':
            if current:
                parts.append(('normal', current))
                current = ""
            # Find closing **
            end = text.find('**', i + 2)
            if end != -1:
                parts.append(('bold', text[i+2:end]))
                i = end + 2
                continue
        
        # Code `text`
        elif text[i] == '`':
            if current:
                parts.append(('normal', current))
                current = ""
            # Find closing `
            end = text.find('`', i + 1)
            if end != -1:
                parts.append(('code', text[i+1:end]))
                i = end + 1
                continue
        
        # Italic *text* (single asterisk, not part of **)
        elif text[i] == '*' and (i == 0 or text[i-1] != '*') and (i == len(text)-1 or text[i+1] != '*'):
            if current:
                parts.append(('normal', current))
                current = ""
            # Find closing *
            end = text.find('*', i + 1)
            if end != -1:
                parts.append(('italic', text[i+1:end]))
                i = end + 1
                continue
        
        current += text[i]
        i += 1
    
    if current:
        parts.append(('normal', current))
    
    # Add formatted runs
    for fmt, txt in parts:
        run = paragraph.add_run(txt)
        if fmt == 'bold':
            run.bold = True
        elif fmt == 'italic':
            run.italic = True
        elif fmt == 'code':
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(231, 76, 60)

if __name__ == "__main__":
    print("Converting PROJECT_DOCUMENTATION.md to Word document...")
    
    try:
        parse_markdown_to_word(
            'PROJECT_DOCUMENTATION.md',
            'PROJECT_DOCUMENTATION.docx'
        )
        print("\n✅ Conversion completed successfully!")
        print("📄 File: PROJECT_DOCUMENTATION.docx")
        
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
